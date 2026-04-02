#!/usr/bin/env python3
"""Word (.docx) → Markdown 変換スクリプト

pandoc が利用可能であれば高精度変換に委譲し、
なければ python-docx で基本的な構造を抽出する。

Usage:
    python3 docx2md.py <input.docx> [--output <output.md>] [--images-dir <dir>]
"""

import argparse
import os
import shutil
import subprocess
import sys
from pathlib import Path


def convert_with_pandoc(input_path: Path, output_path: Path, images_dir: Path) -> bool:
    """pandoc による高精度変換を試みる。成功すれば True を返す。"""
    pandoc = shutil.which("pandoc")
    if not pandoc:
        return False

    cmd = [
        pandoc,
        str(input_path),
        "-t", "markdown",
        "--wrap=none",
        f"--extract-media={images_dir}",
        "-o", str(output_path),
    ]
    try:
        subprocess.run(cmd, check=True, capture_output=True, text=True)
        return True
    except (subprocess.CalledProcessError, FileNotFoundError):
        return False


def convert_with_python_docx(input_path: Path, output_path: Path, images_dir: Path) -> None:
    """python-docx によるフォールバック変換。"""
    try:
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
    except ImportError:
        print("エラー: python-docx が未インストールです。", file=sys.stderr)
        print("  pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(input_path))
    lines: list[str] = []
    image_count = 0

    # 画像を抽出
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1
            images_dir.mkdir(parents=True, exist_ok=True)
            ext = os.path.splitext(rel.target_ref)[1] or ".png"
            img_name = f"image{image_count}{ext}"
            img_path = images_dir / img_name
            with open(img_path, "wb") as f:
                f.write(rel.target_part.blob)

    # 抽出した画像のファイル名マップを作成
    image_filenames: dict[int, str] = {}
    img_idx = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            img_idx += 1
            ext = os.path.splitext(rel.target_ref)[1] or ".png"
            image_filenames[img_idx] = f"image{img_idx}{ext}"

    # 段落を変換
    image_ref_idx = 0
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()

        # 画像の参照を検出（段落にテキストがなくランがある場合）
        has_image = False
        if not text:
            for run in para.runs:
                if run._element.findall(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
                ) or run._element.findall(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict"
                ):
                    has_image = True
                    break
            # inline shape check
            for child in para._element:
                if "drawing" in child.tag or "pict" in child.tag:
                    has_image = True
                    break

        if has_image and image_count > 0:
            image_ref_idx += 1
            if image_ref_idx in image_filenames:
                img_name = image_filenames[image_ref_idx]
                rel_path = f"{images_dir.name}/{img_name}"
                lines.append(f"![image]({rel_path})")
                lines.append("")
            continue

        if not text:
            lines.append("")
            continue

        # 見出し変換
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading", "").strip())
            except ValueError:
                level = 1
            lines.append(f"{'#' * level} {text}")
            lines.append("")
        elif style_name in ("Title",):
            lines.append(f"# {text}")
            lines.append("")
        elif style_name.startswith("List"):
            lines.append(f"- {text}")
        else:
            lines.append(text)
            lines.append("")

    # テーブルを変換
    for table in doc.tables:
        lines.append("")
        for i, row in enumerate(table.rows):
            cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
            if i == 0:
                lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
        lines.append("")

    output_path.write_text("\n".join(lines), encoding="utf-8")


def main():
    parser = argparse.ArgumentParser(description="Word (.docx) → Markdown 変換")
    parser.add_argument("input", help="入力 .docx ファイルパス")
    parser.add_argument("--output", "-o", help="出力 .md ファイルパス（省略時: 入力と同名.md）")
    parser.add_argument("--images-dir", help="画像抽出先ディレクトリ（省略時: <入力ファイル名>_images）")
    args = parser.parse_args()

    input_path = Path(args.input).resolve()
    if not input_path.exists():
        print(f"エラー: ファイルが見つかりません: {input_path}", file=sys.stderr)
        sys.exit(1)
    if input_path.suffix.lower() != ".docx":
        print(f"エラー: .docx ファイルを指定してください: {input_path}", file=sys.stderr)
        sys.exit(1)

    output_path = Path(args.output).resolve() if args.output else input_path.with_suffix(".md")
    images_dir = Path(args.images_dir).resolve() if args.images_dir else input_path.parent / f"{input_path.stem}_images"

    # pandoc を優先
    if convert_with_pandoc(input_path, output_path, images_dir):
        method = "pandoc"
    else:
        print("pandoc が見つかりません。python-docx で変換します。", file=sys.stderr)
        convert_with_python_docx(input_path, output_path, images_dir)
        method = "python-docx"

    print(f"変換完了: {input_path.name} → {output_path.name} (方式: {method})")
    if images_dir.exists() and any(images_dir.iterdir()):
        img_count = len(list(images_dir.iterdir()))
        print(f"  画像: {img_count}件 → {images_dir}")


if __name__ == "__main__":
    main()
